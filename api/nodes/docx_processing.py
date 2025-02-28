import regex as re
import json

from docx import Document
from langchain.schema import AIMessage

from api.utils.state import State



def extract_field(text, field_name):
    pattern =  rf"{field_name}\s*\|\s*(.+)" 
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else None

def extract_isin(text):
    pattern = r"ISIN\s([A-Z]{2}\d{10})"  
    match = re.search(pattern, text)
    return match.group(1) if match else None

def extract_notion(text, field_name):
    pattern = rf"{field_name}\s*\(N\)?\s*\|\s*(.+)"  # Handles (N) and captures value
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else None


def extract_date(text, field_name):
    pattern = rf"{re.escape(field_name)}\s*\|\s*(\d{{1,2}}\s(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s?\d{{4}})"
        
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else None  # Return the full date

def extract_percentage(text, field_name):
    pattern = rf"{re.escape(field_name)}\s*\|\s*([\d\.]+%)(?:\s*of\s*\w+)?"
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else None
    
def read_docx(file_path):
    doc = Document(file_path)
    text_content = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Avoid empty lines
            text_content.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))  # Join cell content

    return "\n".join(text_content)  # Return as a single text block
    


def docx_processing(state: State) -> State:
    """
    This method is used for the docx processing 
    We use here Rule Based processing
    """
   
    print("Started procesing the docx file..............")
   
    file_bytes = state["contents"]
    
    with open("temp.docx", "wb") as f:
        f.write(file_bytes)

    # Extract text from DOCX
    text = read_docx("temp.docx")

    #print(text)
    entities = {}

    # Regex patterns
    entities["counterparty"] = extract_field(text, "Party A")
    entities["ivd"] = extract_date(text, "Initial Valuation Date")
    entities["notional"] =  extract_notion(text, "Notional Amount")
    entities["val_date"] = extract_date(text, "Valuation Date")
    entities["maturity"] = extract_date(text, "Termination Date")
            
    entities["underlying"] = extract_field(text, "Underlying")
    entities["coupon"] = extract_percentage(text, "Coupon (C)")
    entities["barrier"] = extract_percentage(text, "Barrier (B)")
    entities["calender"] = extract_field(text, "Business Day")         

    result= {'file_type': "docx", "entities":entities}

    # Convert dictionary to JSON string
    json_result = json.dumps(result, indent=2)
    
    response_message = AIMessage(content=json_result)
    state["messages"].append(response_message)
         
    print("Finished processing the docx file..............")
    
    return state
    